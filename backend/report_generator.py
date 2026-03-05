from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from groq import Groq
from dotenv import load_dotenv
from pathlib import Path
import os
import datetime

load_dotenv(dotenv_path=Path(__file__).parent / ".env")
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))


def generate_cam_narrative(company_name, analysis, financials, score_result, research):
    prompt = f"""
You are a senior credit analyst at an Indian bank. Write a formal Credit Appraisal Memo (CAM) for:

Company: {company_name}
Financial Analysis: {analysis[:2000]}
Financials: {financials}
Credit Score: {score_result['score']}/100 | Grade: {score_result['grade']}
Decision: {score_result['decision']}
Red Flags: {', '.join(score_result['red_flags']) if score_result['red_flags'] else 'None'}
Web Research Risk: {research['risk_level']}
Research Summary: {research['summary'][:1000]}

Write these 5 sections in formal Indian banking language:

1. BORROWER BACKGROUND (2-3 sentences about company profile)
2. FINANCIAL ANALYSIS (comment on revenue, profit, debt, ratios)
3. RISK ASSESSMENT (red flags, web research findings, litigation)
4. CREDIT OPINION (your professional assessment)
5. RECOMMENDATION (approve/reject with loan amount, rate, conditions)

Use professional language. Mention amounts in Crores. Be specific and direct.
"""

    response = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    return response.choices[0].message.content


def add_heading(doc, text, level=1, color=None):
    para = doc.add_heading(text, level=level)
    if color:
        for run in para.runs:
            run.font.color.rgb = RGBColor(*color)
    return para


def add_table_row(table, label, value, highlight=False):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = str(value)
    if highlight:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.bold = True


def generate_cam_report(company_name, analysis, financials, score_result, research) -> str:
    doc = Document()

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CREDIT APPRAISAL MEMORANDUM (CAM)")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Confidential | {datetime.date.today().strftime('%d %B %Y')}")

    doc.add_paragraph()

    # ── Basic Info Table ──
    add_heading(doc, "1. LOAN APPLICATION SUMMARY", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Details"

    info = [
        ("Borrower Name", company_name),
        ("Date of Appraisal", datetime.date.today().strftime("%d-%m-%Y")),
        ("Credit Score", f"{score_result['score']} / 100"),
        ("Grade", score_result['grade']),
        ("Decision", score_result['decision']),
        ("Suggested Loan Limit", f"INR {score_result['suggested_loan_limit_crore']} Crore"),
        ("Indicative Rate", score_result['interest_rate']),
        ("Web Research Risk", research['risk_level']),
    ]

    for label, value in info:
        highlight = label in ["Decision", "Grade"]
        add_table_row(table, label, value, highlight)

    doc.add_paragraph()

    # ── Financials Table ──
    add_heading(doc, "2. KEY FINANCIALS", level=2)
    fin_table = doc.add_table(rows=1, cols=2)
    fin_table.style = "Table Grid"
    fin_table.rows[0].cells[0].text = "Metric"
    fin_table.rows[0].cells[1].text = "Value (INR Crore)"

    for key, val in financials.items():
        add_table_row(fin_table, key, val)

    doc.add_paragraph()

    # ── Red Flags ──
    add_heading(doc, "3. RED FLAGS & RISK SIGNALS", level=2)
    if score_result['red_flags']:
        for flag in score_result['red_flags']:
            p = doc.add_paragraph(f"⚠ {flag}", style="List Bullet")
    else:
        doc.add_paragraph("No significant red flags identified.")

    doc.add_paragraph()

    # ── Web Research ──
    add_heading(doc, "4. SECONDARY RESEARCH FINDINGS", level=2)
    doc.add_paragraph(f"Risk Level: {research['risk_level']}")
    doc.add_paragraph(research['summary'][:1500])

    doc.add_paragraph()

    # ── LLM Narrative ──
    add_heading(doc, "5. DETAILED CREDIT ASSESSMENT", level=2)
    narrative = generate_cam_narrative(
        company_name, analysis, financials, score_result, research
    )
    doc.add_paragraph(narrative)

    doc.add_paragraph()

    # ── Final Decision Box ──
    add_heading(doc, "6. FINAL RECOMMENDATION", level=2)
    decision_para = doc.add_paragraph()
    decision_run = decision_para.add_run(
        f"DECISION: {score_result['decision']}  |  "
        f"LIMIT: INR {score_result['suggested_loan_limit_crore']} Crore  |  "
        f"RATE: {score_result['interest_rate']}"
    )
    decision_run.bold = True
    decision_run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph(f"Explanation: {score_result['explanation']}")

    doc.add_paragraph()
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Prepared by: Intelli-Credit AI Engine")
    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}")

    # Save
    filename = f"CAM_{company_name.replace(' ', '_')}_{datetime.date.today()}.docx"
    filepath = Path(__file__).parent / filename
    doc.save(filepath)

    return str(filepath)